"""Tests for converters — REGISTRY completeness, actual conversion, on_progress."""

import os
import sys
import tempfile
import threading
from pathlib import Path

# Ensure project root is in path
sys.path.insert(0, str(Path(__file__).parent.parent))

import unittest
from converters import REGISTRY, get_supported_conversions, convert


# ── Sample file generators ──────────────────────────────────────────

def _make_sample_pdf(path: str):
    """Create a minimal 2-page PDF with text and a table."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet

    styles = getSampleStyleSheet()
    story = [
        Paragraph("Hello PDF", styles["Title"]),
        Paragraph("This is page 1 content.", styles["BodyText"]),
        Spacer(1, 20),
        Paragraph("Page 2 content here.", styles["BodyText"]),
        Table(
            [["A", "B"], ["1", "2"]],
            style=TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.grey)]),
        ),
    ]
    SimpleDocTemplate(path, pagesize=A4).build(story)


def _make_sample_xlsx(path: str):
    """Create a minimal Excel with 2 sheets."""
    from openpyxl import Workbook
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sheet1"
    ws1["A1"] = "Name"
    ws1["B1"] = "Value"
    ws1["A2"] = "foo"
    ws1["B2"] = "42"
    ws2 = wb.create_sheet("Sheet2")
    ws2["A1"] = "X"
    ws2["A2"] = "Y"
    wb.save(path)


def _make_sample_docx(path: str):
    """Create a minimal Word document with heading, paragraph, and table."""
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("Some paragraph text.")
    table = doc.add_table(rows=2, cols=2, style="Table Grid")
    table.cell(0, 0).text = "Col1"
    table.cell(0, 1).text = "Col2"
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    doc.save(path)


def _make_sample_md(path: str):
    """Create a minimal Markdown file."""
    Path(path).write_text(
        "# Hello\n\nThis is **bold** and `code`.\n\n## Section 2\n\n- item 1\n- item 2\n",
        encoding="utf-8",
    )


def _make_sample_zip(path: str):
    """Create a minimal ZIP with an HTML file inside."""
    import zipfile
    html = "<html><body><h1>Test</h1><p>Hello world</p></body></html>"
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("test.html", html.encode("utf-8"))


SAMPLE_MAKERS = {
    "pdf": _make_sample_pdf,
    "xlsx": _make_sample_xlsx,
    "docx": _make_sample_docx,
    "md": _make_sample_md,
    "zip": _make_sample_zip,
}

# Conversions that are expected to work
EXPECTED_CONVERSIONS = [
    ("pdf", "xlsx"),
    ("pdf", "docx"),
    ("pdf", "md"),
    ("xlsx", "pdf"),
    ("xlsx", "docx"),
    ("xlsx", "md"),
    ("docx", "pdf"),
    ("docx", "xlsx"),
    ("docx", "md"),
    ("md", "pdf"),
    ("md", "xlsx"),
    ("md", "docx"),
    ("zip", "md"),
]


class TestRegistry(unittest.TestCase):
    """Verify REGISTRY contains all expected conversions."""

    def test_registry_completeness(self):
        supported = set(get_supported_conversions())
        expected = set(EXPECTED_CONVERSIONS)
        self.assertEqual(
            supported, expected,
            f"Missing: {expected - supported}, Extra: {supported - expected}",
        )

    def test_registry_count(self):
        self.assertEqual(len(REGISTRY), 13)


class TestConversions(unittest.TestCase):
    """Run all 13 conversions and verify output exists and is non-empty."""

    def _run_conversion(self, from_ext, to_ext):
        with tempfile.TemporaryDirectory() as tmpdir:
            src = os.path.join(tmpdir, f"sample.{from_ext}")
            SAMPLE_MAKERS[from_ext](src)
            dst = os.path.join(tmpdir, f"output.{to_ext}")

            convert(src, dst, from_ext, to_ext)

            self.assertTrue(
                os.path.isfile(dst),
                f"Output file not created: {dst}",
            )
            self.assertGreater(
                os.path.getsize(dst), 0,
                f"Output file is empty: {dst}",
            )

    def test_all_conversions(self):
        for from_ext, to_ext in EXPECTED_CONVERSIONS:
            with self.subTest(f"{from_ext} -> {to_ext}"):
                self._run_conversion(from_ext, to_ext)

    def test_unsupported_conversion_raises(self):
        with self.assertRaises(ValueError):
            convert("in.pdf", "out.zip", "pdf", "zip")


class TestOnProgress(unittest.TestCase):
    """Verify on_progress callback receives real granular progress."""

    def test_on_progress_pdf_to_md(self):
        updates = []
        with tempfile.TemporaryDirectory() as tmpdir:
            src = os.path.join(tmpdir, "sample.pdf")
            _make_sample_pdf(src)
            dst = os.path.join(tmpdir, "out.md")

            def cb(done, total):
                updates.append((done, total))

            convert(src, dst, "pdf", "md", on_progress=cb)

        self.assertGreater(len(updates), 0, f"Expected many progress updates, got {len(updates)}")
        # Last update should have done == total
        last_done, last_total = updates[-1]
        self.assertEqual(last_done, last_total)

    def test_on_progress_xlsx_to_md(self):
        updates = []
        with tempfile.TemporaryDirectory() as tmpdir:
            src = os.path.join(tmpdir, "sample.xlsx")
            _make_sample_xlsx(src)
            dst = os.path.join(tmpdir, "out.md")

            def cb(done, total):
                updates.append((done, total))

            convert(src, dst, "xlsx", "md", on_progress=cb)

        self.assertGreater(len(updates), 0)
        last_done, last_total = updates[-1]
        self.assertEqual(last_done, last_total)

    def test_on_progress_docx_to_md(self):
        updates = []
        with tempfile.TemporaryDirectory() as tmpdir:
            src = os.path.join(tmpdir, "sample.docx")
            _make_sample_docx(src)
            dst = os.path.join(tmpdir, "out.md")

            def cb(done, total):
                updates.append((done, total))

            convert(src, dst, "docx", "md", on_progress=cb)

        self.assertGreater(len(updates), 0)
        last_done, last_total = updates[-1]
        self.assertEqual(last_done, last_total)


if __name__ == "__main__":
    unittest.main()