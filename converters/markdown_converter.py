import html
import os
import re
from pathlib import Path


def _register_pdf_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fonts_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    for font_path in (fonts_dir / "msyh.ttc", fonts_dir / "simhei.ttf", fonts_dir / "simsun.ttc"):
        if font_path.exists():
            try:
                pdfmetrics.registerFont(TTFont("FileConverterCJK", str(font_path)))
                return "FileConverterCJK"
            except Exception:
                pass
    return "Helvetica"


def _markdown_to_pdf(md_path: str, pdf_path: str):
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "MarkdownBody", parent=styles["BodyText"], fontName=font_name,
        fontSize=10, leading=16, alignment=TA_LEFT, spaceAfter=6,
    )
    heading1 = ParagraphStyle(
        "MarkdownHeading1", parent=body, fontSize=18, leading=24,
        spaceBefore=8, spaceAfter=10,
    )
    heading2 = ParagraphStyle(
        "MarkdownHeading2", parent=body, fontSize=14, leading=20,
        spaceBefore=6, spaceAfter=8,
    )

    story = []
    text = Path(md_path).read_text(encoding="utf-8")
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 4))
            continue
        if stripped.startswith("|---") or (stripped.startswith("|") and "|" in stripped[1:]):
            continue

        style = body
        if stripped.startswith("# "):
            style, stripped = heading1, stripped[2:].strip()
        elif stripped.startswith("## "):
            style, stripped = heading2, stripped[3:].strip()
        elif stripped.startswith(("- ", "* ")):
            stripped = "&#8226; " + stripped[2:].strip()

        escaped = html.escape(stripped)
        escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
        escaped = re.sub(r"`(.+?)`", r"<font name='Courier'>\1</font>", escaped)
        story.append(Paragraph(escaped, style))

    document = SimpleDocTemplate(
        pdf_path, pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
    )
    document.build(story)


class MarkdownConverter:

    @staticmethod
    def to_pdf(md_path: str, pdf_path: str):
        _markdown_to_pdf(md_path, pdf_path)

    @staticmethod
    def to_excel(md_path: str, xlsx_path: str):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Markdown Content"

        text = Path(md_path).read_text(encoding="utf-8")
        row_idx = 1

        for line in text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                if not any("---" in c for c in cells):
                    for col_idx, cell in enumerate(cells, start=1):
                        ws.cell(row=row_idx, column=col_idx, value=cell)
                    row_idx += 1
            elif stripped:
                ws.cell(row=row_idx, column=1, value=stripped)
                row_idx += 1

        wb.save(xlsx_path)

    @staticmethod
    def to_word(md_path: str, docx_path: str):
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        text = Path(md_path).read_text(encoding="utf-8")

        for line in text.split("\n"):
            stripped = line.strip()
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif "|" in stripped and stripped.startswith("|"):
                continue
            elif stripped.startswith("|---"):
                continue
            elif stripped:
                p = doc.add_paragraph(stripped)
                for run in p.runs:
                    run.font.size = Pt(10)

        doc.save(docx_path)