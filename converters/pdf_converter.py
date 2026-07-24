import pdfplumber
from openpyxl import Workbook
from pathlib import Path
from typing import Callable, Optional


class PdfConverter:

    @staticmethod
    def to_excel(pdf_path: str, excel_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            wb = Workbook()
            wb.remove(wb.active)
            for i, page in enumerate(pdf.pages, start=1):
                ws = wb.create_sheet(title=f"Page {i}")
                row_idx = 1

                found_tables = page.find_tables()
                table_bboxes = [t.bbox for t in found_tables]

                def not_in_table(obj):
                    if obj["object_type"] == "char":
                        for bbox in table_bboxes:
                            if (bbox[0] <= obj["x0"] and bbox[2] >= obj["x1"]
                                    and bbox[1] <= obj["top"] and bbox[3] >= obj["bottom"]):
                                return False
                    return True

                text = page.filter(not_in_table).extract_text()
                if text:
                    for line in text.split("\n"):
                        ws.cell(row=row_idx, column=1, value=line)
                        row_idx += 1
                    row_idx += 1

                for table in [t.extract() for t in found_tables]:
                    for row in table:
                        for col_idx, cell in enumerate(row, start=1):
                            ws.cell(row=row_idx, column=col_idx, value=cell if cell else "")
                        row_idx += 1
                    row_idx += 1

                if on_progress:
                    on_progress(i, total_pages)

            wb.save(excel_path)

    @staticmethod
    def to_word(pdf_path: str, docx_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            for idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text:
                    p = doc.add_paragraph(text)
                    for run in p.runs:
                        run.font.size = Pt(10)

                tables = [t.extract() for t in page.find_tables()]
                for table_data in tables:
                    rows_count = len(table_data)
                    cols_count = max(len(r) for r in table_data) if table_data else 0
                    if rows_count == 0 or cols_count == 0:
                        continue
                    table = doc.add_table(rows=rows_count, cols=cols_count, style="Table Grid")
                    for i, row_data in enumerate(table_data):
                        for j, cell in enumerate(row_data):
                            if j < cols_count:
                                table.cell(i, j).text = str(cell) if cell else ""

                doc.add_paragraph()

                if on_progress:
                    on_progress(idx, total_pages)

        doc.save(docx_path)

    @staticmethod
    def to_markdown(pdf_path: str, md_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        lines = []
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages, start=1):
                lines.append(f"## Page {i}\n")
                text = page.extract_text()
                if text:
                    lines.append(text)
                    lines.append("")

                tables = [t.extract() for t in page.find_tables()]
                for table_data in tables:
                    if not table_data:
                        continue
                    header = table_data[0]
                    lines.append("| " + " | ".join(str(h) if h else "" for h in header) + " |")
                    lines.append("|" + "|".join("---" for _ in header) + "|")
                    for row in table_data[1:]:
                        lines.append("| " + " | ".join(str(c) if c else "" for c in row) + " |")
                    lines.append("")

                if on_progress:
                    on_progress(i, total_pages)

        Path(md_path).write_text("\n".join(lines), encoding="utf-8")