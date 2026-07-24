from pathlib import Path
from openpyxl import load_workbook
from typing import Callable, Optional
from .pdf_export import docx_to_pdf


class ExcelConverter:

    @staticmethod
    def to_pdf(xlsx_path: str, pdf_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        from docx import Document

        wb = load_workbook(xlsx_path)
        sheets = wb.worksheets
        total_sheets = len(sheets)
        doc = Document()

        for idx, ws in enumerate(sheets, start=1):
            doc.add_heading(ws.title, level=2)
            table_data = []
            for row in ws.iter_rows(values_only=True):
                table_data.append([str(c) if c else "" for c in row])

            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]), style="Table Grid")
                for i, row_data in enumerate(table_data):
                    for j, cell_val in enumerate(row_data):
                        if j < len(table_data[0]):
                            table.cell(i, j).text = cell_val

            doc.add_paragraph()

            if on_progress:
                # Phase 1: generating docx (half of total)
                on_progress(idx, total_sheets * 2)

        temp_docx = str(Path(pdf_path).with_suffix(".docx"))
        doc.save(temp_docx)

        if on_progress:
            on_progress(total_sheets, total_sheets * 2)

        docx_to_pdf(temp_docx, pdf_path, "Excel -> PDF requires MS Word (Windows) or LibreOffice (Linux/Mac)")

        if on_progress:
            on_progress(total_sheets * 2, total_sheets * 2)

        Path(temp_docx).unlink(missing_ok=True)

    @staticmethod
    def to_word(xlsx_path: str, docx_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        from docx import Document
        from docx.shared import Inches, Pt

        wb = load_workbook(xlsx_path)
        sheets = wb.worksheets
        doc = Document()

        for idx, ws in enumerate(sheets, start=1):
            doc.add_heading(ws.title, level=2)
            table_data = []
            column_widths = []
            for row in ws.iter_rows(values_only=True):
                table_data.append([str(c) if c else "" for c in row])
                if not column_widths:
                    column_widths = [Inches(1.5)] * len(row)

            if table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]), style="Table Grid")
                for i, row_data in enumerate(table_data):
                    for j, cell_val in enumerate(row_data):
                        if j < len(table_data[0]):
                            table.cell(i, j).text = cell_val
                doc.add_paragraph()

            if on_progress:
                on_progress(idx, len(sheets))

        doc.save(docx_path)

    @staticmethod
    def to_markdown(xlsx_path: str, md_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        wb = load_workbook(xlsx_path)
        sheets = wb.worksheets
        lines = []

        for idx, ws in enumerate(sheets, start=1):
            lines.append(f"## {ws.title}\n")
            table_data = []
            for row in ws.iter_rows(values_only=True):
                table_data.append([str(c) if c else "" for c in row])

            if table_data:
                header = table_data[0]
                lines.append("| " + " | ".join(header) + " |")
                lines.append("|" + "|".join("---" for _ in header) + "|")
                for row in table_data[1:]:
                    lines.append("| " + " | ".join(row) + " |")
            else:
                lines.append("(This sheet is empty)")
            lines.append("")

            if on_progress:
                on_progress(idx, len(sheets))

        Path(md_path).write_text("\n".join(lines), encoding="utf-8")