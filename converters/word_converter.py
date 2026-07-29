from pathlib import Path
from docx import Document
from typing import Callable, Optional
from .pdf_export import docx_to_pdf


class WordConverter:

    @staticmethod
    def to_pdf(docx_path: str, pdf_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        if on_progress:
            on_progress(0, 2)
        docx_to_pdf(docx_path, pdf_path, "Word -> PDF requires MS Word (Windows) or LibreOffice (Linux/Mac)")
        if on_progress:
            on_progress(2, 2)

    @staticmethod
    def to_excel(docx_path: str, xlsx_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        from openpyxl import Workbook

        doc = Document(docx_path)
        wb = Workbook()
        ws = wb.active
        ws.title = "Document Content"

        # Count total items for progress
        total_items = max(len(doc.paragraphs) + len(doc.tables), 1)
        done = 0

        row_idx = 1
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                ws.cell(row=row_idx, column=1, value=text)
                row_idx += 1
            done += 1
            if on_progress:
                on_progress(done, total_items)

        for i, table in enumerate(doc.tables):
            row_idx += 1
            ws.cell(row=row_idx, column=1, value=f"[Table {i+1}]")
            row_idx += 1
            for row in table.rows:
                for j, cell in enumerate(row.cells):
                    ws.cell(row=row_idx, column=j + 1, value=cell.text)
                row_idx += 1
            done += 1
            if on_progress:
                on_progress(done, total_items)

        wb.save(xlsx_path)

    @staticmethod
    def to_markdown(docx_path: str, md_path: str, on_progress: Optional[Callable[[int, int], None]] = None):
        doc = Document(docx_path)
        lines = []

        # Count total items for progress
        total_items = max(len(doc.paragraphs) + len(doc.tables), 1)
        done = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                done += 1
                if on_progress:
                    on_progress(done, total_items)
                continue
            style = para.style.name
            if "Heading 1" in style:
                lines.append(f"# {text}")
            elif "Heading 2" in style:
                lines.append(f"## {text}")
            elif "Heading 3" in style:
                lines.append(f"### {text}")
            else:
                lines.append(text)
            lines.append("")
            done += 1
            if on_progress:
                on_progress(done, total_items)

        for i, table in enumerate(doc.tables):
            lines.append(f"**Table {i+1}**\n")
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            if table_data:
                header = table_data[0]
                lines.append("| " + " | ".join(header) + " |")
                lines.append("|" + "|".join("---" for _ in header) + "|")
                for row in table_data[1:]:
                    lines.append("| " + " | ".join(row) + " |")
            lines.append("")
            done += 1
            if on_progress:
                on_progress(done, total_items)

        Path(md_path).write_text("\n".join(lines), encoding="utf-8")