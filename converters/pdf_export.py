import subprocess
import sys
import shutil
from pathlib import Path


def _find_libreoffice():
    """Find LibreOffice executable across platforms."""
    if sys.platform == "darwin":
        mac_paths = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice.bin",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        ]
        for p in mac_paths:
            if Path(p).exists():
                return p
    return shutil.which("libreoffice") or shutil.which("soffice")


def _docx_to_pdf_with_reportlab(docx_path: str, pdf_path: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from .markdown_converter import _register_pdf_font
    from docx import Document

    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "DocxBody", parent=styles["BodyText"], fontName=font_name,
        fontSize=10, leading=16, spaceAfter=6,
    )
    heading = ParagraphStyle(
        "DocxHeading", parent=body, fontSize=14, leading=20,
        spaceBefore=6, spaceAfter=8,
    )
    story = []
    document = Document(docx_path)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            story.append(Spacer(1, 4))
            continue
        style = heading if paragraph.style.name.startswith("Heading") else body
        story.append(Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style))

    for docx_table in document.tables:
        rows = []
        for row in docx_table.rows:
            rows.append([
                Paragraph(
                    cell.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"),
                    body,
                )
                for cell in row.cells
            ])
        if rows:
            table = Table(rows, repeatRows=1)
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.extend([table, Spacer(1, 8)])

    output = SimpleDocTemplate(
        pdf_path, pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
    )
    output.build(story)


def docx_to_pdf(docx_path: str, pdf_path: str, error_message: str):
    """Convert .docx to PDF with Office when available, otherwise ReportLab."""
    try:
        if sys.platform == "win32":
            try:
                import win32com.client

                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                wd = word.Documents.Open(str(Path(docx_path).resolve()))
                wd.SaveAs(str(Path(pdf_path).resolve()), FileFormat=17)
                wd.Close()
                word.Quit()
            except Exception:
                _docx_to_pdf_with_reportlab(docx_path, pdf_path)
        else:
            lo_exe = _find_libreoffice()
            if lo_exe is None:
                _docx_to_pdf_with_reportlab(docx_path, pdf_path)
                return
            try:
                subprocess.run(
                    [lo_exe, "--headless", "--convert-to", "pdf",
                     "--outdir", str(Path(pdf_path).parent.resolve()),
                     str(Path(docx_path).resolve())],
                    check=True, capture_output=True,
                )
            except Exception:
                _docx_to_pdf_with_reportlab(docx_path, pdf_path)
    except Exception as exc:
        raise RuntimeError(f"{error_message}: {exc}") from exc
